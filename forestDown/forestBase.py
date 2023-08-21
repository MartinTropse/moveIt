import os
import re
import arcpy
from arcpy import env
import math
from thefuzz import fuzz
import win32com.client
import docx
import time


def forestAreamap(prjPath, prjCode, xDim=1.1, yDim=1.1, dpi=300):
    if bool(re.search("(?i)[a-z]$", prjCode)):
        prjCode2 = prjCode[:-1]

    # Loop that goes through the download folders and identifies the geodatabases.
    gdbPaths = []
    for bDir, sDir, mFiles in os.walk(prjPath):
        for folder in sDir:
            try:
                if (hasattr(re.search(f"{prjCode}.+\.gdb$", folder), "group")):
                    gdbPaths.append(os.path.join(bDir, folder))
                elif (hasattr(re.search(f"{prjCode2}.+\.gdb$", folder), "group")):
                    gdbPaths.append(os.path.join(bDir, folder))
            except Exception as e:
                print(e)

    layerPaths = []

    for aGdb in gdbPaths:
        arcpy.env.workspace = aGdb
        datasets = arcpy.ListDatasets(feature_type="feature")
        datasets = [""] + datasets if datasets is not None else []
        for ds in datasets:
            for fc in arcpy.ListFeatureClasses(feature_dataset=ds):
                layerPaths.append(os.path.join(arcpy.env.workspace, ds, fc))

    # Identify specific layers.
    AreaRgx = "(?i).+\.gdb.+projekt(om| om|_om).+"
    BuffRgx = "(?i).+\.gdb.+buff(.+1|1)"

    AreaGdb = [i for i in layerPaths if re.search(AreaRgx, i)][0]
    BuffGdb = [i for i in layerPaths if re.search(BuffRgx, i)][0]
    projAreaList = [AreaGdb, BuffGdb]

    aprx_path = os.path.join(prjPath, "Blank.aprx")
    aprx = arcpy.mp.ArcGISProject(aprx_path)
    m = aprx.listMaps("Map")[0]

    xmax = 0
    xmin = 1000000000
    ymax = 0
    ymin = 1000000000

    for gdbLayer in projAreaList:
        gdbPath, layer = os.path.split(gdbLayer)
        fldPath = os.path.split(gdbPath)
        lyrxExp = layer + ".lyrx"
        tempLayer = arcpy.MakeFeatureLayer_management(gdbLayer, layer)
        arcpy.env.workspace = fldPath[0]
        arcpy.management.SaveToLayerFile(layer, lyrxExp)
        lyrxLayer = arcpy.mp.LayerFile(os.path.join(fldPath[0], lyrxExp))
        m.addLayer(lyrxLayer)
        desc = arcpy.Describe(tempLayer)
        extent = desc.extent
        if extent.XMax > xmax:
            xmax = extent.XMax
        if extent.XMin < xmin:
            xmin = extent.XMin
        if extent.YMax > ymax:
            ymax = extent.YMax
        if extent.YMin < ymin:
            ymin = extent.YMin

    BuffGdb = os.path.normpath(BuffGdb)
    buffName = BuffGdb.split(os.sep)[len(BuffGdb.split(os.sep)) - 1]

    AreaGdb = os.path.normpath(AreaGdb)
    areaName = AreaGdb.split(os.sep)[len(AreaGdb.split(os.sep)) - 1]

    buffLyr = m.listLayers(buffName)[0]
    buffLyr.name = "1 km buffertzon"

    projLyr = m.listLayers(areaName)[0]
    projLyr.name = "Projektområde"

    proj_cim = projLyr.getDefinition("V3")
    buff_cim = buffLyr.getDefinition("V3")

    proj_cim.renderer.symbol.symbol.symbolLayers[0].color.values = [255, 0, 197, 100]
    proj_cim.renderer.symbol.symbol.symbolLayers[0].width = 1.0
    # proj_cim.renderer.symbol.symbol.symbolLayers[1].color.values = [255, 0, 197, 100]
    proj_cim.renderer.symbol.symbol.symbolLayers[1].enable = "False"
    projLyr.setDefinition(proj_cim)

    buff_cim.renderer.symbol.symbol.symbolLayers[0].width = 1.5
    buff_cim.renderer.symbol.symbol.symbolLayers[0].color.values = [45, 45, 44, 100]
    buff_cim.renderer.symbol.symbol.symbolLayers[1].enable = "False"
    buffLyr.setDefinition(buff_cim)

    # Set the extent of the exported map
    extent.YMax = ymax + (ymax - ymin) * yDim
    extent.XMax = xmax + (xmax - xmin) * xDim
    extent.YMin = ymin - (ymax - ymin) * yDim
    extent.XMin = xmin - (xmax - xmin) * xDim

    layout = aprx.listLayouts("Kartmall - Rapportanpassad fyrkantig")[0]
    map_frame = layout.listElements("MAPFRAME_ELEMENT", "braveNewFrame")[0]

    lay_legend = layout.listElements("legend_element", "*")[0]
    lay_legend.items[0]

    map_frame.camera.setExtent(extent)
    output_path = os.path.join(prjPath, (prjCode + "_SurveyAreaMap.jpg"))

    layout.exportToJPEG(output_path, resolution=dpi)
    aprx.saveACopy(os.path.join(prjPath, prjCode + "_SurveyAreaMap.aprx"))
    print(f"Copied map image and ArcGIS-projekt file to\n{output_path}")
    return output_path


def Direction(angle):
  if((angle >= 150 and angle <= 180) or (angle >= -180 and angle <= -150)):
      return "Syd"
  if((angle <= -121 and angle >= -149)):
      return "Sydväst"
  if(angle <= -60 and angle >= -120):
      return "Väst"
  if((angle <= 30 and angle >= 0) or (angle >= -30 and angle <= 0)):
      return "Norr"
  if(angle <= -31 and angle >= -59):
      return "Nordväst"
  if(angle >= 60 and angle <= 120):
      return "Öst"
  if(angle >= 31 and angle <= 59):
      return "Nordöst"
  if(angle >= 121 and angle <= 149):
      return "Sydöst"


month_namesDict = {
    'January': 'januari',
    'February': 'februari',
    'March': 'mars',
    'April': 'april',
    'May': 'maj',
    'June': 'juni',
    'July': 'juli',
    'August': 'augusti',
    'September': 'september',
    'October': 'oktober',
    'November': 'november',
    'December': 'december'
}
#Dictionary to convert interge to numeric str
num_dict = {"1":"en",
            "2":"två",
            "3":"tre",
            "4":"fyra",
            "5":"fem",
            "6":"sex",
            "7":"sju",
            "8":"åtta",
            "9":"nio"}

def roundup1000(x):
    return int(math.ceil(x / 1000.0)) * 1000
def roundup10(x):
    return int(math.ceil(x / 10.0)) * 10

fieldPersDict = {"Jof":"Johan Frölinghaus", "JohanF":"Johan Frölinghaus", "Johan Frölinghaus":"JF", "persvs":"Per-Erik Svahn","persva":"Per-Erik Svahn", "pesva":"Per-Erik Svahn", "Annika Rastén":"Annika Rastén",
                 "Josefina":"Josefina Pehrson", "Josefina Pehrson":"Josefina Pehrson","JP":"Josefina Pehrson", "Frida Nettelbladt":"Frida Nettelbladt","Frida Nettelbladt":"FN","Sam Keith":"Samuel Keith","Samuel Keith":"Samuel Keith",
                 "SK":"Samuel Keith", "OKK":"Olle Kvarnbäck", "Olle Kvarnbäck":"Olle Kvarnbäck", "Per Österman":"Per Österman", "PÖ":"Per Österman","Hans Nibon":"Hans Nibon", "HN":"Hans Nibon", "Samuel Holdar":"Samuel Holdar", "Samuel Holdar":"SH",
                 "Erik Edvardsson":"Erik Edvardsson", "EE":"Erik Edvardsson","Gabriel Säll":"Gabriel Säll", "GS":"Gabriel Säll", "Isabella Rasmusson Honnér":"Isabella Rasmusson Honnér", "IRH":"Isabella Rasmusson Honnér","Rickard Gustavsson":"Rickard Gustavsson",
                 "RG":"Rickard Gustavsson","Marie Björklund":"Marie Björklund","MB":"Marie Björklund", "Jakob Sörensen":"Jakob Sörensen", "JS":"Jakob Sörensen", "MH":"Mårten Hjernquist", "Milad Avalinejad-Bandari":"Milad Avalinejad-Bandari", "MAB":"Milad Avalinejad-Bandari",
                 "Anders Tranberg":"Anders Tranberg", "AT":"Anders Tranberg", "Arvid Löf":"Arvid Löf", "AL":"Arvid Löf","Edwin Sahlin":"Edwin Sahlin", "ES":"Edwin Sahlin","Håkan Andersson":"Håkan Andersson","HA":"Håkan Andersson",
                 "Ofir Svensson":"Ofir Svensson", "OS":"Ofir Svensson","Mova Hebert":"Mova Hebert","MH":"Mova Hebert","Petter Andersson":"Petter Andersson","PA":"Petter Andersson","Eric Bjuringer":"Eric Bjuringer","EB":"Eric Bjuringer",
                 "Emil Lundahl":"Emil Lundahl","EL":"Emil Lundahl", "Christopher Gullander":"Christopher Gullander","CG":"Christopher Gullander", "Marcus Bergström":"Marcus Bergström","MB":"Marcus Bergström"}


def trueFieldPers(fieldpersNames):
    for fPos, fName in enumerate(fieldpersNames):
        bestScore = 0
        for kName in fieldPersDict.keys():
            nameScore = fuzz.ratio(fName.lower(),kName.lower())
            if len(fName.split(" ")) >= 3: #This section usually does not run, but handles cases when multiple names has been written on a observation
                nameScore = fuzz.partial_ratio(fName.lower(),kName.lower())
                if nameScore > bestScore:
                    bestScore = nameScore
                    bestName = fieldPersDict[kName]
            elif nameScore > bestScore:
                bestScore = nameScore
                bestName = fieldPersDict[kName]
        if len(fName) < 4:
            if bestScore > 66:
                fieldpersNames[fPos] = bestName
            else:
                print(f"Could not find a matching name for: {fName}")
        if len(fName) >= 4:
            if bestScore > 80:
                fieldpersNames[fPos] = bestName
            elif len(fName.split(" ")) == 2:
                continue
            else:
                print(f"Could not find a matching name for: {fName}")
    return set(fieldpersNames), fieldpersNames

def remvEmptyParagraphs(wordPath):
    """Function that goes through the word-document and removes empty sections. Parameters: The function expects a path to the word file to be modified"""
    doc = docx.Document(wordPath)
    paraRange = range(0, len(doc.paragraphs)-1)
    removeList = []
    #Loop that identifies the position of headers followed by empty text paragraphs
    for parPos in paraRange:
        if re.search("(?i)heading.+", doc.paragraphs[parPos].style.name):
            #This if statement checks that (1) The text following header is empty, (2) is classified as "Normal" text (3) that it is not in last position of the doc object.
            if len(doc.paragraphs[parPos+1].text) == 0 and doc.paragraphs[parPos+1].style.name == "Normal" and parPos != len(doc.paragraphs):
                removeList.append(parPos)
                removeList.append(parPos+1)
    #Removes the paragraphs and following whitespace
    for rmPos in removeList:
        print(f"\nDeleted this paragraph: {doc.paragraphs[rmPos].text}")
        print(rmPos)
        time.sleep(1)
        doc.paragraphs[rmPos].clear()
        p1 = doc.paragraphs[rmPos+1]._element
        p1.getparent().remove(p1)
        p0 = doc.paragraphs[rmPos]._element
        p0.getparent().remove(p0)
    doc.save(wordPath)


def braveNewTOC(wordPath):
    word = win32com.client.DispatchEx("Word.Application")
    doc = word.Documents.Open(wordPath)
    doc.TablesOfContents(1).Update()
    doc.Close(SaveChanges=True)
    word.Quit()
